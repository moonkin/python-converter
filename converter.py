import configparser
from datetime import datetime
import docx
import os


class Document:
    """Class for parsing documents and modifying their contents."""

    def __init__(self, fname, config_fname='basic.ini'):
        self.fname = fname
        self.format = Format(config_fname)
        self.paragraphs = []

    def __getitem__(self, i):
        return self.paragraphs[i]

    def get_modified_text(self):
        """Return a list of div tags and paragraphs wrapped in p tags."""
        result = [self.format.div_tag]
        for par in self.paragraphs:
            string = par.get_string()
            if string:
                if par.align == 'right':
                    p_tag_right = self.format.p_tag[:2] + " align='right'" + \
                                  self.format.p_tag[2:]
                    wr_par = p_tag_right + string + self.format.p_end_tag
                else:
                    wr_par = self.format.p_tag + string + self.format.p_end_tag
                result.append(wr_par)
            else:
                result.append(self.format.div_end_tag + self.format.div_tag)
        result.append(self.format.div_end_tag)

        return result


class TxtDocument(Document):
    """Subclass for .txt documents."""

    def __init__(self, fname):
        super().__init__(fname)
        with open(self.fname, 'r', encoding='cp1251') as f:
            for par in f.readlines():
                self.paragraphs.append(TxtPar(par))

    def __str__(self):
        return ''.join([par.get_text for par in self.paragraphs])


class DocxDocument(Document):
    """Subclass for .docx documents."""

    def __init__(self, fname):
        super().__init__(fname)
        for par in docx.Document(self.fname).paragraphs:
            self.paragraphs.append(DocxPar(par))

    def __str__(self):
        return '\n'.join(list(map(str, self.paragraphs)))


class Par:
    """Class for paragraphs of text."""

    def __init__(self, par):
        self.paragraph = par
        self.align = 'justify'


class TxtPar(Par):
    """Subclass for paragraphs of text from .txt documents."""

    def __str__(self):
        return self.paragraph

    def get_string(self):
        return self.paragraph.strip()


class DocxPar(Par):
    """Subclass for paragraphs of text from .docx documents."""

    def __init__(self, par):
        super().__init__(par)
        if 'RIGHT' in str(self.paragraph.alignment):
            self.align = 'right'

    def __str__(self):
        return self.paragraph.text

    def get_string(self):
        """Return the paragraph with basic formatting added."""
        runs = self.paragraph.runs
        pairs_of_runs = list(zip(runs[:], runs[1:]))
        if runs:
            cur = runs[0]
            string = DocxPar.add_tags(cur) + cur.text

            for pair_of_runs in pairs_of_runs:
                prev, cur = pair_of_runs
                tags = DocxPar.add_tags(cur, prev)
                string += tags + cur.text

            string += DocxPar.add_tags(None, cur)
            return string

        return ''

    @staticmethod
    def add_tags(cur, prev=None):
        """Compare formatting of two runs and return proper html tags."""
        tags = ''
        if cur:
            if cur.bold and (not prev or not prev.bold):
                tags += '<b>'
            if cur.italic and (not prev or not prev.italic):
                tags += '<i>'
            if cur.underline and (not prev or not prev.underline):
                tags += '<u>'
        if prev:
            if (not cur or not cur.underline) and prev.underline:
                tags += '</u>'
            if (not cur or not cur.italic) and prev.italic:
                tags += '</i>'
            if (not cur or not cur.bold) and prev.bold:
                tags += '</b>'
        return tags


class Format:
    """Class that reads text formatting settings from the config file."""

    def __init__(self, config_fname):
        config = configparser.ConfigParser()
        config.read(config_fname)
        section = config['DEFAULT']
        ind = self.get_indent(section)
        w = self.get_width(section)
        p_mar = self.get_margins('p', section).rstrip()
        div_mar = self.get_margins('div', section)

        self.div_tag = f"<div align='justify' style='{div_mar}{ind}{w}'>"
        self.p_tag = f"<p style='{p_mar}'>" if p_mar else '<p>'
        self.div_end_tag = '</div>'
        self.p_end_tag = '</p>'

    def get_width(self, section):
        width = int(section['width'])
        units = section['width-units']
        if width != 100:
            return 'width: {0}{1};'.format(width, units)
        return ''

    def get_indent(self, section):
        indent = float(section['text-indent'])
        units = section['text-indent-units']
        if indent:
            return 'text-indent: {0}{1};'.format(indent, units)
        return ''

    def get_margins(self, tag, section):
        top = float(section[f'{tag}-margin-top'])
        right = float(section[f'{tag}-margin-right'])
        bottom = float(section[f'{tag}-margin-bottom'])
        left = float(section[f'{tag}-margin-left'])
        units = section[f'{tag}-margin-units']
                                                                    # symbols
        if not all([top, right, bottom, left]):                     #   0
            return ''
        elif top == right == bottom == left:                        #   14
            print(top, right, left, right)
            print(not right, not left, not top, not bottom)
            return 'margin: {1}{0}; '.format(units, top)
        elif top and not all([right, bottom, left]):                #   18
            return 'margin-top: {1}{0}; '.format(units, top)
        elif top == bottom and right == left:                       #   19
            return 'margin: {1}{0} {2}{0}; '.format(units, top, right)
        elif left and not all([right, top, bottom]):                #   19
            return 'margin-right: {1}{0}; '.format(units, left)
        elif right and not all([left, top, bottom]):                #   20
            return 'margin-left: {1}{0}; '.format(units, right)
        elif bottom and not all([right, left, top]):                #   21
            return 'margin-bottom: {1}{0}; '.format(units, bottom)
        elif top != bottom and right == left:                       #   26
            return 'margin: {1}{0} {2}{0} {3}{0}; '.format(units, top,
                                                           right, bottom)
        else:                                                       #   32
            return 'margin: {1}{0} {2}{0} {3}{0} {4}{0}; '.format(units,
                                                                  top, right,
                                                                  bottom, left)


if __name__ == "__main__":

    timestamp = datetime.today().strftime('_%d_%m_%Y')
    idir = os.path.abspath(os.getcwd())
    odir = os.path.abspath(os.path.join(idir, 'Converted' + timestamp))

    if not os.path.exists(odir):
        os.mkdir(odir)

    for fname in os.listdir(idir):

        file_name, file_format = os.path.splitext(fname)
        if file_format == '.txt':
            document = TxtDocument(fname)
        elif file_format in ['.docx', '.doc']:
            document = DocxDocument(fname)
        else:
            continue

        html_document = document.get_modified_text()

        fpath_txt = os.path.join(odir, file_name + '.txt')
        fpath_html = fpath_txt.replace('.txt', '.html')

        for fpath in [fpath_txt, fpath_html]:
            with open(fpath, 'w', encoding='cp1251') as f:
                for line in html_document:
                    f.write(line)

